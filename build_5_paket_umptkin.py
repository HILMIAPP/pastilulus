from __future__ import annotations

import random
import re
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"G:\BISMILLAH 100JT\New folder\lolosujian")
SOURCE_DIR = Path(r"G:\BISMILLAH 100JT\New folder\SOAL UMPTKIN\New folder")
SOAL_PATH = SOURCE_DIR / "Soal_HOTS_UM-PTKIN_2026.docx"
PEMBAHASAN_PATH = SOURCE_DIR / "Pembahasan_HOTS_UM-PTKIN_2026.docx"
OUT_DIR = ROOT / "hasil_5_paket_umptkin_hots"
SEPARATE_DIR = OUT_DIR / "paket_terpisah"


@dataclass
class Question:
    original_number: int
    section: str
    duration: str
    lines: list[str]
    answer: str
    explanation: str = ""


def nonempty_paragraphs(path: Path) -> list[str]:
    doc = Document(str(path))
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def parse_questions() -> list[Question]:
    paragraphs = nonempty_paragraphs(SOAL_PATH)
    questions: list[Question] = []
    section = ""
    duration = ""
    block: list[str] = []

    def flush() -> None:
        nonlocal block
        if not block:
            return
        match = re.match(r"^(\d+)\.\s*(.*)$", block[0])
        if not match:
            block = []
            return
        number = int(match.group(1))
        answer_line = next((line for line in block if line.startswith("Jawaban:")), "")
        answer = answer_line.split(":", 1)[1].strip() if ":" in answer_line else ""
        lines = block[:]
        lines[0] = match.group(2)
        lines = [line for line in lines if not line.startswith("Jawaban:")]
        questions.append(Question(number, section, duration, lines, answer))
        block = []

    for text in paragraphs:
        if text == "KUNCI JAWABAN":
            flush()
            break
        if text.startswith("SUB TES"):
            flush()
            section = text
            duration = ""
            continue
        if section and re.match(r"^\d+\s+Soal", text):
            duration = text
            continue
        if re.match(r"^\d+\.\s+", text):
            flush()
            block = [text]
            continue
        if block:
            block.append(text)

    flush()
    return questions


def parse_explanations() -> list[str]:
    paragraphs = nonempty_paragraphs(PEMBAHASAN_PATH)
    explanations: list[str] = []
    for idx, text in enumerate(paragraphs):
        if text == "Pembahasan:" and idx + 1 < len(paragraphs):
            explanations.append(paragraphs[idx + 1])
    return explanations


def apply_source_corrections(questions: list[Question]) -> None:
    # Source question 17 has answer C, while the worked solution proves B.
    for question in questions:
        if question.original_number == 17:
            question.answer = "B"
            question.explanation = (
                "Selisih antarsuku: 95-100 = -5, 85-95 = -10, "
                "70-85 = -15, 50-70 = -20. Pola selisih berkurang 5, "
                "sehingga selisih berikutnya -25 dan -30. Suku ke-6 = "
                "50 - 25 = 25, suku ke-7 = 25 - 30 = -5. Selisih antara "
                "suku ke-7 dan suku ke-6 adalah -5 - 25 = -30. Jadi "
                "jawaban yang benar adalah B."
            )


def package_order(questions: list[Question], package_index: int) -> list[Question]:
    sections: list[str] = []
    for question in questions:
        if question.section not in sections:
            sections.append(question.section)

    result: list[Question] = []
    for section in sections:
        group = [q for q in questions if q.section == section]
        if package_index > 1:
            rng = random.Random(2026 + package_index * 97 + len(group))
            rng.shuffle(group)
        result.extend(group)
    return result


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in [
        ("Title", 20, "0B2545", 0, 6),
        ("Subtitle", 11, "555555", 0, 10),
        ("Heading 1", 16, "2E74B5", 14, 6),
        ("Heading 2", 13, "2E74B5", 10, 4),
        ("Heading 3", 11, "1F4D78", 8, 3),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15


def add_footer(doc: Document, label: str) -> None:
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(label)
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 100, 100)


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def add_package_header(doc: Document, package_label: str, subtitle: str) -> None:
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(package_label).bold = True

    sub = doc.add_paragraph(style="Subtitle")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(subtitle)


def add_question(doc: Document, number: int, question: Question, include_answer: bool = True) -> None:
    first = doc.add_paragraph()
    first.paragraph_format.keep_with_next = True
    first.add_run(f"{number}. ").bold = True
    first.add_run(question.lines[0])

    for line in question.lines[1:]:
        paragraph = doc.add_paragraph()
        if re.match(r"^[A-E]\.\s+", line):
            paragraph.paragraph_format.left_indent = Inches(0.22)
        paragraph.add_run(line)

    if include_answer:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.22)
        run = paragraph.add_run(f"Jawaban: {question.answer}")
        run.bold = True
        run.font.color.rgb = RGBColor(31, 77, 120)


def add_answer_key(doc: Document, ordered: list[Question]) -> None:
    doc.add_paragraph("KUNCI JAWABAN", style="Heading 2")
    chunk_size = 20
    for start in range(0, len(ordered), chunk_size):
        chunk = ordered[start : start + chunk_size]
        text = "   ".join(f"{start + idx + 1}. {q.answer}" for idx, q in enumerate(chunk))
        doc.add_paragraph(text)


def build_soal_doc(packages: list[list[Question]]) -> Path:
    doc = Document()
    setup_styles(doc)
    add_footer(doc, "5 Paket HOTS UM-PTKIN 2026")

    for idx, ordered in enumerate(packages, start=1):
        if idx > 1:
            add_page_break(doc)
        add_package_header(
            doc,
            f"PAKET {idx} - SOAL HOTS UM-PTKIN 2026",
            "121 soal setara berdasarkan komposisi subtes UM-PTKIN",
        )

        current_section = None
        for number, question in enumerate(ordered, start=1):
            if question.section != current_section:
                current_section = question.section
                doc.add_paragraph(question.section, style="Heading 1")
                doc.add_paragraph(question.duration, style="Heading 3")
            add_question(doc, number, question)

        add_answer_key(doc, ordered)

    out = OUT_DIR / "Soal_HOTS_UMPTKIN_2026_5_Paket.docx"
    doc.save(str(out))
    return out


def build_pembahasan_doc(packages: list[list[Question]]) -> Path:
    doc = Document()
    setup_styles(doc)
    add_footer(doc, "Pembahasan 5 Paket HOTS UM-PTKIN 2026")

    for idx, ordered in enumerate(packages, start=1):
        if idx > 1:
            add_page_break(doc)
        add_package_header(
            doc,
            f"PAKET {idx} - PEMBAHASAN HOTS UM-PTKIN 2026",
            "Kunci dan pembahasan mengikuti urutan soal pada paket terkait",
        )

        current_section = None
        for number, question in enumerate(ordered, start=1):
            if question.section != current_section:
                current_section = question.section
                doc.add_paragraph(question.section.replace("SUB TES", "PEMBAHASAN SUB TES"), style="Heading 1")
                doc.add_paragraph(question.duration, style="Heading 3")

            heading = doc.add_paragraph(style="Heading 3")
            heading.add_run(f"{number}. Jawaban: {question.answer}").bold = True
            body = doc.add_paragraph()
            body.add_run(question.explanation)

    out = OUT_DIR / "Pembahasan_HOTS_UMPTKIN_2026_5_Paket.docx"
    doc.save(str(out))
    return out


def build_single_soal_doc(ordered: list[Question], package_index: int) -> Path:
    doc = Document()
    setup_styles(doc)
    add_footer(doc, f"Paket {package_index} HOTS UM-PTKIN 2026")
    add_package_header(
        doc,
        f"PAKET {package_index} - SOAL HOTS UM-PTKIN 2026",
        "121 soal setara berdasarkan komposisi subtes UM-PTKIN",
    )

    current_section = None
    for number, question in enumerate(ordered, start=1):
        if question.section != current_section:
            current_section = question.section
            doc.add_paragraph(question.section, style="Heading 1")
            doc.add_paragraph(question.duration, style="Heading 3")
        add_question(doc, number, question)

    add_answer_key(doc, ordered)
    out = SEPARATE_DIR / f"Paket_{package_index}_Soal_HOTS_UMPTKIN_2026.docx"
    doc.save(str(out))
    return out


def build_single_pembahasan_doc(ordered: list[Question], package_index: int) -> Path:
    doc = Document()
    setup_styles(doc)
    add_footer(doc, f"Pembahasan Paket {package_index} HOTS UM-PTKIN 2026")
    add_package_header(
        doc,
        f"PAKET {package_index} - PEMBAHASAN HOTS UM-PTKIN 2026",
        "Kunci dan pembahasan mengikuti urutan soal pada paket terkait",
    )

    current_section = None
    for number, question in enumerate(ordered, start=1):
        if question.section != current_section:
            current_section = question.section
            doc.add_paragraph(question.section.replace("SUB TES", "PEMBAHASAN SUB TES"), style="Heading 1")
            doc.add_paragraph(question.duration, style="Heading 3")

        heading = doc.add_paragraph(style="Heading 3")
        heading.add_run(f"{number}. Jawaban: {question.answer}").bold = True
        doc.add_paragraph(question.explanation)

    out = SEPARATE_DIR / f"Paket_{package_index}_Pembahasan_HOTS_UMPTKIN_2026.docx"
    doc.save(str(out))
    return out


def build_separate_docs(packages: list[list[Question]]) -> list[Path]:
    SEPARATE_DIR.mkdir(exist_ok=True)
    outputs: list[Path] = []
    for idx, ordered in enumerate(packages, start=1):
        outputs.append(build_single_soal_doc(ordered, idx))
        outputs.append(build_single_pembahasan_doc(ordered, idx))
    return outputs


def main() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    questions = parse_questions()
    explanations = parse_explanations()
    if len(questions) != 121:
        raise ValueError(f"Expected 121 questions, found {len(questions)}")
    if len(explanations) != len(questions):
        raise ValueError(f"Expected {len(questions)} explanations, found {len(explanations)}")

    for question, explanation in zip(questions, explanations):
        question.explanation = explanation
    apply_source_corrections(questions)

    packages = [package_order(deepcopy(questions), idx) for idx in range(1, 6)]
    soal_out = build_soal_doc(packages)
    pembahasan_out = build_pembahasan_doc(packages)
    separate_out = build_separate_docs(packages)

    print(soal_out)
    print(pembahasan_out)
    for path in separate_out:
        print(path)


if __name__ == "__main__":
    main()
