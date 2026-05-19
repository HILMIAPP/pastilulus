from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from textwrap import dedent

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"G:\BISMILLAH 100JT\New folder\lolosujian")
OUT_DIR = ROOT / "materi_bimbel_umptkin_a_z"
ASSET_DIR = OUT_DIR / "gambar"


@dataclass
class Exercise:
    prompt: str
    choices: list[str]
    answer: str
    discussion: str


@dataclass
class Chapter:
    title: str
    opening: str
    concepts: list[tuple[str, str]]
    patterns: list[str]
    example: Exercise
    drills: list[Exercise]


@dataclass
class Module:
    code: str
    title: str
    subtitle: str
    image: str
    overview: str
    chapters: list[Chapter]
    glossary: list[tuple[str, str]]


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf",
    ]
    for path in candidates:
        try:
            return ImageFont.truetype(path, size)
        except OSError:
            pass
    return ImageFont.load_default()


def draw_wrapped(draw: ImageDraw.ImageDraw, xy: tuple[int, int], text: str, fnt, width: int, fill: str) -> int:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        trial = (current + " " + word).strip()
        if draw.textbbox((0, 0), trial, font=fnt)[2] <= width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    x, y = xy
    line_height = draw.textbbox((0, 0), "Ag", font=fnt)[3] + 8
    for line in lines:
        draw.text((x, y), line, font=fnt, fill=fill)
        y += line_height
    return y


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str) -> None:
    draw.line([start, end], fill=color, width=4)
    x1, y1 = start
    x2, y2 = end
    if x2 >= x1:
        head = [(x2, y2), (x2 - 14, y2 - 8), (x2 - 14, y2 + 8)]
    else:
        head = [(x2, y2), (x2 + 14, y2 - 8), (x2 + 14, y2 + 8)]
    draw.polygon(head, fill=color)


def make_roadmap_image(path: Path) -> None:
    img = Image.new("RGB", (1600, 900), "#F7FAFC")
    d = ImageDraw.Draw(img)
    d.rounded_rectangle((40, 40, 1560, 860), radius=28, fill="#FFFFFF", outline="#B7C9D9", width=3)
    d.text((90, 80), "Peta Belajar UM-PTKIN dari A sampai Z", font=font(46, True), fill="#0B2545")
    d.text((92, 145), "Bangun konsep, latih pola, ukur kelemahan, lalu ulangi dengan target.", font=font(26), fill="#345")
    boxes = [
        ("Penalaran Akademik", "logika, analogi, pola, inferensi"),
        ("Penalaran Matematika", "model, hitung, grafik, peluang"),
        ("Literasi Membaca", "gagasan, sikap penulis, koherensi"),
        ("Literasi Ajaran Islam", "Quran-Hadits, fikih, SKI, moderasi"),
    ]
    colors = ["#E8F1FF", "#EAF7EA", "#FFF4D8", "#F1E9FF"]
    xs = [100, 480, 860, 1240]
    for x, (title, desc), c in zip(xs, boxes, colors):
        d.rounded_rectangle((x, 250, x + 260, 470), radius=20, fill=c, outline="#9FB3C8", width=3)
        d.text((x + 24, 280), title, font=font(25, True), fill="#17324D")
        draw_wrapped(d, (x + 24, 325), desc, font(22), 210, "#334155")
    for i in range(3):
        arrow(d, (xs[i] + 270, 360), (xs[i + 1] - 15, 360), "#4C6F91")
    stages = [
        ("1. Diagnosis", "cek kemampuan awal"),
        ("2. Konsep", "pahami rumus dan kaidah"),
        ("3. Latihan", "naikkan level soal"),
        ("4. Evaluasi", "catat salah dan ulangi"),
    ]
    for i, (a, b) in enumerate(stages):
        x = 140 + i * 355
        d.ellipse((x, 610, x + 95, 705), fill="#0B5CAD")
        d.text((x + 33, 636), str(i + 1), font=font(34, True), fill="white")
        d.text((x + 115, 615), a, font=font(27, True), fill="#0B2545")
        d.text((x + 115, 655), b, font=font(22), fill="#334155")
    img.save(path)


def make_logic_image(path: Path) -> None:
    img = Image.new("RGB", (1500, 850), "#FFFFFF")
    d = ImageDraw.Draw(img)
    d.text((70, 55), "Alur Menjawab Soal Penalaran", font=font(42, True), fill="#0B2545")
    steps = [
        ("Baca Stimulus", "temukan data eksplisit"),
        ("Tandai Relasi", "sebab-akibat, syarat, kuantor"),
        ("Uji Opsi", "cari yang paling wajib benar"),
        ("Eliminasi", "buang ekstrem dan di luar teks"),
    ]
    y = 180
    for i, (t, s) in enumerate(steps):
        x = 90 + i * 350
        d.rounded_rectangle((x, y, x + 280, y + 180), radius=20, fill="#E8F1FF", outline="#4C6F91", width=3)
        d.text((x + 25, y + 30), t, font=font(27, True), fill="#17324D")
        draw_wrapped(d, (x + 25, y + 78), s, font(22), 220, "#334155")
        if i < len(steps) - 1:
            arrow(d, (x + 290, y + 90), (x + 340, y + 90), "#4C6F91")
    d.rounded_rectangle((160, 520, 1340, 730), radius=18, fill="#F8FAFC", outline="#CBD5E1", width=2)
    d.text((205, 555), "Kunci HOTS:", font=font(30, True), fill="#0B2545")
    draw_wrapped(
        d,
        (205, 605),
        "Jawaban benar bukan yang terdengar bagus, melainkan yang paling kuat didukung oleh data, aturan logika, dan batasan soal.",
        font(25),
        1060,
        "#334155",
    )
    img.save(path)


def make_math_image(path: Path) -> None:
    img = Image.new("RGB", (1500, 850), "#FFFFFF")
    d = ImageDraw.Draw(img)
    d.text((70, 55), "Dari Cerita ke Model Matematika", font=font(42, True), fill="#0B2545")
    d.line((170, 650, 1320, 650), fill="#334155", width=3)
    d.line((250, 700, 250, 190), fill="#334155", width=3)
    pts = []
    for x in range(-10, 11):
        px = 250 + (x + 10) * 45
        y = 0.045 * (x - 1) ** 2 + 1.2
        py = int(650 - y * 80)
        pts.append((px, py))
    d.line(pts, fill="#0B5CAD", width=5)
    d.text((920, 210), "Contoh grafik kuadrat", font=font(30, True), fill="#0B2545")
    draw_wrapped(d, (920, 260), "Titik puncak membantu menjawab nilai maksimum/minimum tanpa mencoba semua nilai.", font(23), 430, "#334155")
    d.rounded_rectangle((90, 120, 550, 230), radius=18, fill="#EAF7EA", outline="#6AA36F", width=3)
    d.text((120, 148), "Teks soal -> variabel -> persamaan -> jawaban", font=font(24, True), fill="#17324D")
    img.save(path)


def make_reading_image(path: Path) -> None:
    img = Image.new("RGB", (1500, 850), "#FFFFFF")
    d = ImageDraw.Draw(img)
    d.text((70, 55), "Anatomi Paragraf Literasi", font=font(42, True), fill="#0B2545")
    d.rounded_rectangle((110, 170, 930, 650), radius=20, fill="#F8FAFC", outline="#CBD5E1", width=3)
    blocks = [
        ("Kalimat utama", "#E8F1FF", "gagasan pusat yang mengikat isi paragraf"),
        ("Data/ilustrasi", "#EAF7EA", "bukti, contoh, angka, atau peristiwa pendukung"),
        ("Simpulan tersirat", "#FFF4D8", "makna yang harus ditarik dari hubungan antarbagian"),
    ]
    y = 220
    for title, color, desc in blocks:
        d.rounded_rectangle((160, y, 880, y + 105), radius=14, fill=color, outline="#9FB3C8", width=2)
        d.text((190, y + 18), title, font=font(28, True), fill="#17324D")
        d.text((190, y + 58), desc, font=font(22), fill="#334155")
        y += 130
    d.rounded_rectangle((1010, 210, 1390, 610), radius=20, fill="#F1E9FF", outline="#9C7BC2", width=3)
    d.text((1050, 250), "Tiga Pertanyaan", font=font(30, True), fill="#17324D")
    for i, line in enumerate(["Apa yang dibahas?", "Bagaimana buktinya?", "Apa konsekuensinya?"]):
        d.text((1060, 320 + i * 75), f"{i + 1}. {line}", font=font(25), fill="#334155")
    img.save(path)


def make_islam_image(path: Path) -> None:
    img = Image.new("RGB", (1500, 850), "#FFFFFF")
    d = ImageDraw.Draw(img)
    d.text((70, 55), "Peta Materi Literasi Ajaran Islam", font=font(42, True), fill="#0B2545")
    center = (750, 430)
    d.ellipse((610, 290, 890, 570), fill="#E8F1FF", outline="#4C6F91", width=4)
    d.text((662, 392), "Ajaran Islam", font=font(32, True), fill="#0B2545")
    items = [
        ("Quran-Hadits", 180, 190, "#EAF7EA"),
        ("Aqidah-Akhlak", 1080, 190, "#FFF4D8"),
        ("Fikih", 180, 590, "#F1E9FF"),
        ("SKI dan Moderasi", 1050, 590, "#F8FAFC"),
    ]
    for title, x, y, color in items:
        d.rounded_rectangle((x, y, x + 270, y + 110), radius=18, fill=color, outline="#9FB3C8", width=3)
        d.text((x + 25, y + 35), title, font=font(27, True), fill="#17324D")
        arrow(d, center, (x + 135, y + 55), "#4C6F91")
    d.text((450, 730), "HOTS menuntut penerapan nilai, bukan hafalan istilah semata.", font=font(28, True), fill="#334155")
    img.save(path)


def create_images() -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    images = {
        "roadmap": ASSET_DIR / "roadmap_umptkin.png",
        "logic": ASSET_DIR / "alur_penalaran.png",
        "math": ASSET_DIR / "model_matematika.png",
        "reading": ASSET_DIR / "anatomi_literasi.png",
        "islam": ASSET_DIR / "peta_ajaran_islam.png",
    }
    make_roadmap_image(images["roadmap"])
    make_logic_image(images["logic"])
    make_math_image(images["math"])
    make_reading_image(images["reading"])
    make_islam_image(images["islam"])
    return images


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.columns[i].width = Inches(widths[i])
        set_cell_text(table.rows[0].cells[i], header, True)
        shade_cell(table.rows[0].cells[i], "E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].width = Inches(widths[i])
            set_cell_text(cells[i], text)
    doc.add_paragraph()


def setup_doc(doc: Document, module_title: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Title", 22, "0B2545", 0, 8),
        ("Subtitle", 12, "555555", 0, 12),
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        st = styles[style_name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.25

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(module_title)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 100, 100)


def para(doc: Document, text: str, style: str | None = None, bold_start: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    if bold_start and text.startswith(bold_start):
        r = p.add_run(bold_start)
        r.bold = True
        p.add_run(text[len(bold_start) :])
    else:
        p.add_run(text)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_exercise(doc: Document, exercise: Exercise, number: int, with_discussion: bool = False) -> None:
    para(doc, f"{number}. {exercise.prompt}")
    for choice in exercise.choices:
        doc.add_paragraph(choice, style="List Bullet")
    if with_discussion:
        para(doc, f"Jawaban: {exercise.answer}", bold_start="Jawaban:")
        para(doc, f"Pembahasan: {exercise.discussion}", bold_start="Pembahasan:")


def extra_topics(code: str) -> list[tuple[str, str, str]]:
    data = {
        "00": [
            ("Analisis kebutuhan siswa", "Petakan skor awal, kebiasaan belajar, kemampuan baca, dan ketahanan mengerjakan soal panjang.", "Buat tabel nilai awal per subtes dan beri warna pada dua kelemahan terbesar."),
            ("Jadwal 30 hari", "Bagi waktu menjadi blok konsep, drill, review, dan simulasi. Hindari belajar satu subtes saja selama seminggu penuh.", "Tulis jadwal 6 hari belajar + 1 hari evaluasi selama empat pekan."),
            ("Catatan kesalahan", "Setiap salah harus diberi label: konsep, baca, hitung, waktu, atau ragu. Label ini lebih penting daripada sekadar nilai.", "Ambil 10 soal salah terakhir, lalu kelompokkan penyebabnya."),
            ("Strategi tutor", "Tutor perlu menjelaskan cara berpikir, bukan hanya jawaban. Minta siswa menyebut alasan memilih dan alasan menolak opsi lain.", "Saat membahas soal, tanyakan: bukti teksnya di mana? rumusnya dari mana?"),
            ("Simulasi bertahap", "Mulai dari simulasi per bab, lalu per subtes, lalu full tryout. Ketahanan ujian dibangun seperti latihan fisik.", "Coba simulasi 20 menit tanpa membuka catatan, lalu refleksi 5 menit."),
            ("Manajemen panik", "Panik membuat siswa membaca ulang tanpa arah. Ajarkan napas pendek, tandai soal, dan lanjut ke soal yang lebih jelas.", "Latih aturan: jika 90 detik buntu, tandai dan lanjut."),
        ],
        "01": [
            ("Sinonim dan antonim", "Pelajari makna dalam konteks, bukan daftar kata terpisah. Kata yang sama bisa berubah rasa makna menurut kalimat.", "Buat 5 kalimat dengan kata 'tajam' dalam makna berbeda."),
            ("Analogi kompleks", "Tentukan relasi sebelum melihat opsi: pencipta-karya, alat-fungsi, tempat-aktivitas, sebab-akibat, atau tingkatan.", "Cari pasangan analogi untuk: editor : naskah."),
            ("Silogisme", "Gunakan diagram atau simbol sederhana agar tidak tertipu kuantor semua, sebagian, dan tidak semua.", "Ubah kalimat 'Sebagian A bukan B' menjadi contoh konkret."),
            ("Pernyataan cukup-perlu", "Syarat cukup membawa akibat, syarat perlu harus ada agar akibat mungkin terjadi.", "Buat contoh: memiliki tiket adalah syarat perlu untuk masuk bioskop."),
            ("Konsistensi data", "Jika soal memberi beberapa pernyataan, cek apakah ada yang bertentangan atau tidak cukup informasi.", "Buat tabel benar/salah untuk tiga pernyataan sederhana."),
            ("Pola visual", "Amati jumlah, posisi, arah, warna, dan urutan. Jangan langsung menebak dari bentuk yang paling mirip.", "Gambar 3 kotak dengan pola rotasi 90 derajat."),
            ("Argumen kritis", "Bedakan klaim, bukti, asumsi, dan simpulan. Opsi yang melemahkan biasanya menyerang asumsi.", "Ambil satu iklan dan tulis klaim serta buktinya."),
        ],
        "02": [
            ("Operasi bilangan", "Kuasai pecahan, persen, perbandingan, dan estimasi. Banyak soal panjang selesai jika angka disederhanakan.", "Ubah 12,5 persen menjadi pecahan sederhana."),
            ("Perbandingan", "Perbandingan dapat berupa bagian:bagian atau bagian:total. Kesalahan membaca bentuk ini sangat sering terjadi.", "Jika A:B=2:3 dan total 50, tentukan A."),
            ("Aritmetika sosial", "Diskon bertingkat tidak dijumlahkan langsung. Gunakan perkalian persentase sisa.", "Hitung harga setelah diskon 20 persen lalu 10 persen."),
            ("Barisan dan deret", "Cek beda, rasio, beda bertingkat, dan pola campuran. Tuliskan minimal lima suku sebelum menyimpulkan.", "Lanjutkan 3, 7, 15, 31, ..."),
            ("Fungsi dan grafik", "Grafik adalah cerita visual. Titik potong, puncak, dan kemiringan punya makna dalam konteks.", "Jelaskan arti gradien dalam grafik jarak-waktu."),
            ("Geometri ruang", "Buat koordinat sederhana untuk kubus atau balok jika gambar sulit dibayangkan.", "Tentukan diagonal ruang balok 3,4,12."),
            ("Peluang", "Pastikan ruang sampel setara. Jika tidak setara, gunakan kasus atau pohon peluang.", "Buat ruang sampel pelemparan dua koin."),
            ("Statistika", "Mean, median, dan modus menjawab kebutuhan berbeda. Data berpencilan biasanya lebih cocok dengan median.", "Cari median dari 2, 4, 4, 5, 30."),
        ],
        "03": [
            ("Skimming", "Gunakan skimming untuk membaca kerangka teks: topik, arah argumen, dan paragraf penting.", "Baca satu teks 400 kata dalam 2 menit, tulis satu kalimat inti."),
            ("Scanning", "Scanning mencari informasi spesifik. Mata bergerak ke angka, nama, istilah, atau kata kunci pertanyaan.", "Cari semua angka dalam artikel pendek dan jelaskan maksudnya."),
            ("Gagasan utama", "Gagasan utama lebih luas daripada contoh. Jika opsi hanya menyebut detail kecil, biasanya salah.", "Pilih kalimat yang bisa menjadi payung untuk tiga kalimat lain."),
            ("Inferensi", "Inferensi harus dekat dengan teks. Hindari jawaban yang benar secara umum tetapi tidak didukung bacaan.", "Tulis satu simpulan dari paragraf berita, lalu garis bawahi buktinya."),
            ("Sikap penulis", "Sikap terlihat dari diksi: ironis, prihatin, optimistis, kritis, atau netral.", "Cari tiga kata evaluatif dalam editorial."),
            ("Vocabulary in context", "Makna kata asing dapat ditebak dari lawan kata, contoh, atau sebab-akibat di sekitar kalimat.", "Tebak makna 'scarce' dari kalimat tentang air yang sulit diperoleh."),
            ("Dhamir Arab", "Kata ganti dalam teks Arab harus dikembalikan ke isim yang tepat agar isi bacaan tidak keliru.", "Tentukan rujukan kata 'huwa' dalam kalimat pendek buatan sendiri."),
        ],
        "04": [
            ("Tema Quran", "Kelompokkan tema ayat: tauhid, ibadah, akhlak, sosial, ilmu, lingkungan, dan keadilan.", "Ambil satu tema akhlak, tulis contoh penerapan di sekolah."),
            ("Fungsi Hadits", "Hadits menjelaskan Quran, memperinci praktik, dan memberi teladan Nabi dalam kehidupan.", "Beri contoh hadits tentang niat dan kaitkan dengan belajar."),
            ("Aqidah", "Aqidah membentuk cara pandang. Soal HOTS sering meminta sikap yang sesuai iman pada situasi konkret.", "Jelaskan hubungan iman kepada Allah dengan kejujuran."),
            ("Akhlak digital", "Nilai Islam berlaku di ruang digital: tabayyun, menjaga lisan, tidak menyebar aib, dan bertanggung jawab.", "Buat aturan grup kelas yang sesuai akhlak Islam."),
            ("Fikih ibadah", "Pahami syarat, rukun, sunnah, dan pembatal. Bedakan mana inti ibadah dan mana penyempurna.", "Bedakan syarat sah shalat dan rukun shalat."),
            ("Muamalah", "Prinsip transaksi adalah keadilan, kejujuran, kejelasan akad, dan tidak merugikan pihak lain.", "Analisis kasus jual beli barang cacat yang tidak diberitahukan."),
            ("SKI", "Sejarah harus dipahami sebagai sebab-akibat dan kontribusi, bukan tanggal yang terpisah.", "Tulis satu kontribusi ilmuwan Muslim dan dampaknya."),
            ("Moderasi beragama", "Moderasi berarti adil dan seimbang, tetap berprinsip, menghargai perbedaan, dan menolak kekerasan.", "Buat contoh sikap moderat saat berbeda pendapat."),
        ],
    }
    return data[code]


def add_az_appendix(doc: Document, module: Module) -> None:
    doc.add_page_break()
    para(doc, "Lampiran Materi A-Z", "Heading 1")
    para(
        doc,
        "Bagian ini memperluas materi inti. Gunakan sebagai bahan pengayaan, tugas rumah, atau sesi remedial. "
        "Setiap topik ditulis dengan tiga fokus: apa yang harus dikuasai, kesalahan umum yang perlu dihindari, "
        "dan latihan mini yang bisa langsung dikerjakan.",
    )
    rows = []
    for topic, mastery, mini in extra_topics(module.code):
        rows.append([topic, mastery, mini])
    add_table(doc, ["Topik", "Yang Harus Dikuasai", "Latihan Mini"], rows, [1.45, 3.0, 1.95])


def add_study_plan(doc: Document, module: Module) -> None:
    para(doc, "Rencana Belajar 7 Sesi", "Heading 1")
    rows = []
    for i in range(1, 8):
        if i == 1:
            focus = "Diagnosis dan pembacaan peta bab"
        elif i in (2, 3):
            focus = "Konsep inti dan contoh terarah"
        elif i in (4, 5):
            focus = "Drill bertahap dan catatan kesalahan"
        elif i == 6:
            focus = "Simulasi waktu per submateri"
        else:
            focus = "Review, remedial, dan target berikutnya"
        rows.append([f"Sesi {i}", focus, "Tulis skor, jenis salah, dan rencana perbaikan."])
    add_table(doc, ["Sesi", "Fokus", "Output"], rows, [0.9, 3.0, 2.5])


def cover(doc: Document, module: Module, images: dict[str, Path]) -> None:
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(module.title).bold = True
    sub = doc.add_paragraph(style="Subtitle")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(module.subtitle)
    doc.add_picture(str(images[module.image]), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(
        doc,
        "Catatan: modul ini adalah materi bimbel mandiri yang disusun sebagai bahan belajar komprehensif. "
        "Struktur mengikuti area kemampuan UM-PTKIN yang umum dipublikasikan: penalaran akademik, "
        "penalaran matematika, literasi membaca, dan literasi ajaran Islam. Selalu cocokkan kembali "
        "jadwal dan ketentuan teknis dengan laman resmi penyelenggara.",
    )


def write_module(module: Module, images: dict[str, Path]) -> Path:
    doc = Document()
    setup_doc(doc, module.title)
    cover(doc, module, images)
    doc.add_page_break()

    para(doc, "Cara Menggunakan Modul Ini", "Heading 1")
    para(doc, module.overview)
    bullets(
        doc,
        [
            "Baca konsep dasar terlebih dahulu, lalu kerjakan contoh terarah tanpa melihat pembahasan.",
            "Setelah latihan, tulis jenis kesalahan: salah konsep, salah membaca, salah hitung, atau salah strategi.",
            "Ulangi bab yang banyak salahnya sebelum pindah ke simulasi waktu.",
            "Gunakan modul ini bersama bank soal dan catatan kesalahan pribadi agar perkembangan terukur.",
        ],
    )

    para(doc, "Peta Isi Modul", "Heading 1")
    add_table(
        doc,
        ["Bab", "Fokus", "Target Kemampuan"],
        [[str(i + 1), ch.title, ch.opening[:170] + "..."] for i, ch in enumerate(module.chapters)],
        [0.7, 2.2, 3.5],
    )

    for idx, ch in enumerate(module.chapters, start=1):
        doc.add_page_break()
        para(doc, f"Bab {idx}. {ch.title}", "Heading 1")
        para(doc, ch.opening)
        para(doc, "A. Konsep Inti", "Heading 2")
        for concept, explanation in ch.concepts:
            para(doc, concept, "Heading 3")
            para(doc, explanation)
        para(doc, "B. Pola Soal yang Sering Muncul", "Heading 2")
        bullets(doc, ch.patterns)
        para(doc, "C. Langkah Kerja", "Heading 2")
        numbered(
            doc,
            [
                "Tentukan kata kunci dan jenis informasi yang diminta soal.",
                "Ubah informasi menjadi bentuk yang lebih sederhana: tabel, simbol, garis waktu, atau sketsa.",
                "Bandingkan setiap opsi dengan data yang benar-benar tersedia.",
                "Pilih jawaban yang paling kuat, bukan yang paling panjang atau paling menarik.",
            ],
        )
        para(doc, "D. Contoh Terarah", "Heading 2")
        add_exercise(doc, ch.example, 1, with_discussion=True)
        para(doc, "E. Latihan Mandiri", "Heading 2")
        for no, drill in enumerate(ch.drills, start=1):
            add_exercise(doc, drill, no, with_discussion=False)
        para(doc, "Kunci Latihan", "Heading 3")
        for no, drill in enumerate(ch.drills, start=1):
            para(doc, f"{no}. {drill.answer}. {drill.discussion}")

    doc.add_page_break()
    add_az_appendix(doc, module)
    add_study_plan(doc, module)
    para(doc, "Glosarium Cepat", "Heading 1")
    add_table(doc, ["Istilah", "Makna Praktis"], [[a, b] for a, b in module.glossary], [1.8, 4.6])
    para(doc, "Checklist Penguasaan", "Heading 1")
    bullets(
        doc,
        [
            "Saya mampu menjelaskan konsep utama tanpa membaca ulang teks.",
            "Saya mampu mengerjakan contoh soal dengan alasan yang jelas.",
            "Saya punya daftar kesalahan yang harus diperbaiki sebelum simulasi.",
            "Saya mampu menyelesaikan latihan dengan batas waktu yang makin ketat.",
        ],
    )

    out = OUT_DIR / f"{module.code}_{module.title.replace(' ', '_').replace('-', '_')}.docx"
    doc.save(str(out))
    return out


def ex(prompt: str, answer: str, discussion: str, choices: list[str] | None = None) -> Exercise:
    if choices is None:
        choices = [
            "A. Pernyataan pertama saja benar",
            "B. Pernyataan kedua saja benar",
            "C. Kedua pernyataan benar",
            "D. Kedua pernyataan salah",
            "E. Tidak dapat ditentukan",
        ]
    return Exercise(prompt, choices, answer, discussion)


def build_modules() -> list[Module]:
    roadmap_chapters = [
        Chapter(
            "Mengenal Karakter UM-PTKIN",
            "Bab ini menjelaskan cara membaca ujian sebagai sistem kemampuan, bukan sekadar kumpulan soal. Peserta perlu memahami bahwa soal HOTS menguji alasan, hubungan, dan ketepatan mengambil keputusan.",
            [
                ("Apa yang diuji", "UM-PTKIN mengukur kemampuan akademik dan literasi yang dibutuhkan untuk belajar di perguruan tinggi keagamaan Islam negeri. Materi tidak cukup dihafal; peserta harus mampu memakai informasi dalam situasi baru."),
                ("Makna HOTS", "HOTS berarti higher order thinking skills: menganalisis, mengevaluasi, dan menyimpulkan. Pada soal seperti ini, satu kata dalam stimulus bisa mengubah jawaban, sehingga membaca teliti lebih penting daripada menebak cepat."),
                ("Empat area besar", "Kerangka belajar yang aman adalah membagi materi ke dalam penalaran akademik, penalaran matematika, literasi membaca, dan literasi ajaran Islam. Masing-masing punya strategi berbeda, tetapi semuanya menuntut ketelitian."),
            ],
            ["Soal berbasis stimulus panjang", "Opsi yang mirip secara makna", "Pertanyaan yang meminta simpulan paling tepat", "Soal numerik dengan konteks cerita"],
            ex("Seorang siswa selalu benar konsep tetapi sering salah karena tidak membaca kata 'kecuali'. Masalah utamanya adalah ...", "B", "Masalahnya bukan kemampuan konsep, melainkan kontrol membaca instruksi.", ["A. Lemah hafalan", "B. Lemah ketelitian instruksi", "C. Lemah matematika", "D. Tidak menguasai agama", "E. Tidak cocok belajar mandiri"]),
            [
                ex("Jika target utama bulan pertama adalah membangun dasar, kegiatan paling tepat adalah ...", "A", "Bulan pertama sebaiknya fokus konsep dan diagnosis, bukan simulasi penuh setiap hari.", ["A. Diagnosis, konsep, latihan ringan", "B. Tryout penuh tiga kali sehari", "C. Menghafal kunci jawaban", "D. Hanya membaca pembahasan", "E. Menunda latihan sampai hafal semua"]),
                ex("Catatan kesalahan yang baik minimal memuat ...", "C", "Jenis salah dan perbaikan membuat belajar berikutnya lebih terarah."),
            ],
        ),
        Chapter(
            "Metode Belajar Bimbel yang Efektif",
            "Bimbel yang baik bukan hanya memberi soal, tetapi mengatur siklus belajar. Peserta harus tahu kapan membaca teori, kapan latihan, kapan simulasi, dan kapan memperbaiki kesalahan.",
            [
                ("Siklus 4D", "Diagnosis, dalami konsep, drill bertahap, dan debrief. Empat langkah ini mencegah siswa merasa belajar banyak tetapi kesalahan yang sama terus berulang."),
                ("Belajar aktif", "Membaca materi harus disertai kegiatan: menandai kata kunci, membuat peta konsep, menjelaskan ulang dengan bahasa sendiri, dan menulis alasan setiap pilihan jawaban."),
                ("Review berjarak", "Materi yang sudah dikuasai tetap perlu diulang. Review singkat setelah 1 hari, 3 hari, dan 7 hari membantu memindahkan konsep ke memori jangka panjang."),
            ],
            ["Siswa merasa paham saat membaca, tetapi gagal saat diberi soal baru", "Latihan banyak tanpa evaluasi jenis salah", "Belajar satu subtes terlalu lama dan mengabaikan subtes lain"],
            ex("Mengapa pembahasan sebaiknya dibaca setelah mencoba soal?", "C", "Mencoba dulu membuat otak aktif memetakan masalah; pembahasan kemudian mengoreksi proses, bukan menggantikan proses."),
            [
                ex("Kegiatan yang paling menunjukkan belajar aktif adalah ...", "C", "Menjelaskan ulang dan membuat contoh baru memaksa pemahaman, bukan sekadar pengenalan."),
                ex("Fungsi review berjarak adalah ...", "A", "Review berjarak menjaga retensi dan mengurangi lupa."),
            ],
        ),
        Chapter(
            "Strategi Waktu dan Simulasi",
            "Ujian bukan hanya soal bisa atau tidak bisa. Peserta juga harus mengelola waktu, energi, dan urutan pengerjaan agar nilai optimal.",
            [
                ("Tiga putaran pengerjaan", "Putaran pertama untuk soal mudah, putaran kedua untuk soal sedang, putaran ketiga untuk soal yang butuh pemikiran panjang. Cara ini mengurangi risiko kehabisan waktu pada soal yang sebenarnya bisa dijawab."),
                ("Batas waktu mikro", "Setiap jenis soal perlu batas waktu. Jika melewati batas dan belum ada arah, tandai lalu lanjut. Kedisiplinan kecil seperti ini sering menaikkan skor."),
                ("Evaluasi simulasi", "Nilai tryout bukan vonis, melainkan data. Lihat subtes mana yang paling lemah, jenis soal apa yang paling sering salah, dan apakah kesalahan muncul karena konsep atau waktu."),
            ],
            ["Soal jebakan yang menyita waktu", "Soal bacaan panjang di akhir sesi", "Kecenderungan mengganti jawaban benar menjadi salah karena panik"],
            ex("Dalam simulasi, seorang peserta menghabiskan 8 menit pada satu soal. Strategi yang lebih baik adalah ...", "D", "Tandai soal sulit dan lanjut agar soal lain yang lebih mudah tidak hilang peluangnya.", ["A. Tetap sampai selesai", "B. Menghapus jawaban semua soal lain", "C. Membaca ulang materi", "D. Tandai lalu lanjut", "E. Menyerah pada subtes itu"]),
            [
                ex("Tryout paling bermanfaat jika setelahnya peserta ...", "B", "Analisis salah adalah jantung peningkatan skor."),
                ex("Putaran pertama pengerjaan sebaiknya berisi soal ...", "A", "Soal mudah mengamankan skor dan membangun ritme."),
            ],
        ),
    ]

    pa_chapters = [
        Chapter(
            "Logika Pernyataan dan Silogisme",
            "Bab ini membangun dasar penalaran formal: memahami syarat, akibat, kuantor, dan kesimpulan yang sah. Banyak soal HOTS terlihat seperti bacaan biasa, padahal kuncinya ada pada struktur logika.",
            [
                ("Implikasi", "Pernyataan 'jika P maka Q' tidak sama dengan 'jika Q maka P'. Kesalahan membalik hubungan adalah jebakan klasik."),
                ("Modus ponens dan tollens", "Jika P maka Q; P terjadi; maka Q. Jika P maka Q; Q tidak terjadi; maka P tidak terjadi. Dua bentuk ini sering muncul dalam soal kebijakan, sains, atau sosial."),
                ("Kuantor", "Kata semua, sebagian, tidak semua, hanya, dan kecuali harus diperlakukan seperti sinyal logika. 'Tidak semua A adalah B' berarti ada A yang bukan B."),
            ],
            ["Menentukan kesimpulan sah dari dua premis", "Mencari pernyataan yang melemahkan argumen", "Membedakan perlu dan cukup"],
            ex("Jika semua program yang dievaluasi rutin meningkat mutunya, dan Program X dievaluasi rutin, maka ...", "C", "Strukturnya semua P -> Q; X adalah P; maka X adalah Q."),
            [ex("Tidak semua siswa yang rajin membaca mampu menjawab soal inferensi. Makna logisnya adalah ...", "A", "Tidak semua berarti sebagian tidak."), ex("Jika literasi naik maka produktivitas naik. Produktivitas tidak naik. Kesimpulan validnya ...", "B", "Ini modus tollens.")],
        ),
        Chapter(
            "Analogi, Relasi Kata, dan Makna",
            "Soal analogi tidak cukup dijawab dengan asosiasi umum. Peserta harus menemukan hubungan yang paling presisi antara dua kata atau konsep.",
            [
                ("Jenis relasi", "Relasi bisa berupa pencipta-karya, alat-fungsi, bagian-keseluruhan, sebab-akibat, lawan makna, atau tingkatan."),
                ("Presisi hubungan", "Dua opsi bisa tampak benar, tetapi hanya satu yang memiliki pola hubungan sama kuat dengan pasangan awal."),
                ("Makna konotatif", "Kata yang dipakai secara kiasan sering menjadi fokus. Bedakan makna harfiah dari makna sosial, emosional, atau idiomatik."),
            ],
            ["Analogi dua pasangan", "Sinonim dan antonim kontekstual", "Makna kata dalam kalimat"],
            ex("Dokter : Diagnosis = Hakim : ...", "A", "Dokter menghasilkan diagnosis dari pemeriksaan; hakim menghasilkan putusan dari persidangan.", ["A. Putusan", "B. Pengadilan", "C. Saksi", "D. Perkara", "E. Undang-undang"]),
            [ex("Kata 'menyulut' dalam 'isu itu menyulut perdebatan' bermakna ...", "B", "Menyulut dipakai secara kiasan: memicu."), ex("Relasi 'arsitek : rancangan' paling dekat dengan ...", "C", "Keduanya pembuat rancangan/karya konseptual.")],
        ),
        Chapter(
            "Pola Bilangan, Huruf, dan Gambar",
            "Pola menguji kemampuan melihat keteraturan. Peserta perlu mencoba beberapa kemungkinan secara sistematis, bukan menebak dari dua suku pertama saja.",
            [
                ("Selisih bertingkat", "Jika selisih pertama tidak tetap, cek selisih kedua. Banyak barisan HOTS memakai pola +2, +4, +8 atau -3, -6, -12."),
                ("Posisi alfabet", "Ubah huruf menjadi angka agar pola terlihat. A=1 sampai Z=26, lalu cek kenaikan atau penurunan."),
                ("Transformasi gambar", "Untuk gambar, amati rotasi, pencerminan, jumlah elemen, perubahan warna, dan posisi. Buat daftar perubahan dari kiri ke kanan."),
            ],
            ["Barisan aritmetika/geometri tersembunyi", "Gabungan huruf dan angka", "Matriks gambar 3 x 3"],
            ex("2, 6, 18, 54, ... suku berikutnya adalah ...", "D", "Setiap suku dikali 3, sehingga 54 x 3 = 162.", ["A. 72", "B. 108", "C. 144", "D. 162", "E. 216"]),
            [ex("A3, C6, F10, J15, ... berikutnya adalah ...", "C", "Huruf naik +2,+3,+4,+5 dan angka naik +3,+4,+5,+6."), ex("100, 95, 85, 70, 50, ... beda suku ke-7 dan ke-6 adalah ...", "B", "Selisihnya -5,-10,-15,-20,-25,-30.")],
        ),
        Chapter(
            "Membaca Argumen dan Simpulan",
            "Argumen terdiri atas klaim, alasan, bukti, dan asumsi. Soal HOTS sering meminta simpulan, kelemahan, atau pernyataan yang paling memperkuat.",
            [
                ("Klaim utama", "Klaim adalah posisi penulis. Biasanya tampak dari kalimat penilaian atau rekomendasi."),
                ("Asumsi", "Asumsi adalah jembatan yang tidak diucapkan tetapi dibutuhkan agar argumen bekerja."),
                ("Opsi ekstrem", "Kata selalu, pasti, semua, tidak pernah, dan satu-satunya sering membuat opsi terlalu kuat dibanding stimulus."),
            ],
            ["Simpulan paling tepat", "Pernyataan yang melemahkan", "Asumsi tersembunyi"],
            ex("Sebuah kota menambah taman dan setelah itu kualitas udara membaik. Simpulan paling hati-hati adalah ...", "B", "Korelasi tidak otomatis membuktikan satu-satunya sebab.", ["A. Taman pasti satu-satunya penyebab", "B. Penambahan taman mungkin berkontribusi", "C. Kualitas udara tidak berubah", "D. Semua kota harus sama", "E. Data pasti salah"]),
            [ex("Argumen paling kuat jika didukung oleh ...", "A", "Data relevan dan metode jelas memperkuat klaim."), ex("Opsi dengan kata 'semua' harus dicurigai karena ...", "D", "Sering terlalu umum melampaui data.")],
        ),
    ]

    math_chapters = [
        Chapter(
            "Aljabar Dasar sampai Persamaan",
            "Aljabar adalah bahasa untuk mengubah cerita menjadi model. Dalam ujian, kesalahan terbesar sering bukan rumus, melainkan salah menentukan variabel.",
            [
                ("Variabel", "Pilih variabel untuk hal yang ditanya atau hal yang paling mudah dihubungkan dengan data."),
                ("Persamaan linear", "Gunakan keseimbangan: operasi pada ruas kiri harus dilakukan juga pada ruas kanan."),
                ("Faktorisasi", "Faktorisasi mempercepat penyelesaian persamaan kuadrat, terutama jika akar bilangan bulat."),
            ],
            ["Cerita umur dan harga", "Sistem persamaan dua variabel", "Persamaan kuadrat dari konteks luas atau gerak"],
            ex("Jumlah dua bilangan 30 dan selisihnya 8. Bilangan terbesar adalah ...", "C", "Misal x+y=30 dan x-y=8, maka 2x=38, x=19.", ["A. 11", "B. 15", "C. 19", "D. 22", "E. 38"]),
            [ex("Jika 3x+5=20, maka x=...", "A", "3x=15 sehingga x=5."), ex("Harga 2 buku dan 1 pensil Rp17.000, harga 1 buku dan 1 pensil Rp10.000. Harga buku adalah ...", "B", "Kurangkan persamaan, harga buku Rp7.000.")],
        ),
        Chapter(
            "Fungsi, Grafik, dan Optimasi",
            "Fungsi menghubungkan input dan output. Soal HOTS sering meminta membaca makna grafik atau mencari nilai maksimum/minimum.",
            [
                ("Fungsi linear", "Grafik garis menunjukkan perubahan tetap. Gradien menunjukkan laju perubahan."),
                ("Fungsi kuadrat", "Grafik parabola memiliki titik puncak. Untuk ax^2+bx+c, absis puncak adalah -b/(2a)."),
                ("Interpretasi", "Angka dalam fungsi harus dikaitkan kembali dengan konteks: meter, rupiah, waktu, jumlah orang, atau peluang."),
            ],
            ["Titik puncak parabola", "Grafik pertumbuhan", "Komposisi fungsi sederhana"],
            ex("Nilai minimum f(x)=x^2-6x+10 adalah ...", "A", "Puncak di x=3; f(3)=9-18+10=1.", ["A. 1", "B. 3", "C. 6", "D. 9", "E. 10"]),
            [ex("Gradien garis y=2x+3 adalah ...", "A", "Koefisien x adalah gradien."), ex("Jika h(t)=-5t^2+20t+1, tinggi maksimum terjadi saat t=...", "B", "t=-b/2a=-20/(-10)=2.")],
        ),
        Chapter(
            "Geometri dan Pengukuran",
            "Geometri menuntut kemampuan membayangkan bentuk. Gambar bantu sering lebih penting daripada hafalan rumus.",
            [
                ("Bangun datar", "Luas dan keliling harus dibedakan. Perhatikan satuan: cm, cm^2, atau cm^3."),
                ("Pythagoras", "Untuk segitiga siku-siku, kuadrat sisi miring sama dengan jumlah kuadrat dua sisi lainnya."),
                ("Kesebangunan", "Jika dua bangun sebangun, perbandingan sisi bersesuaian sama. Luas berubah menurut kuadrat skala."),
            ],
            ["Jarak titik dalam kubus", "Luas gabungan", "Skala peta dan kesebangunan"],
            ex("Segitiga siku-siku memiliki sisi siku 6 cm dan 8 cm. Sisi miringnya ...", "C", "sqrt(36+64)=sqrt(100)=10.", ["A. 7", "B. 8", "C. 10", "D. 12", "E. 14"]),
            [ex("Skala 1:200. Panjang gambar 4 cm berarti panjang sebenarnya ...", "D", "4 x 200 = 800 cm = 8 m."), ex("Luas persegi sisi 9 cm adalah ...", "B", "9 x 9 = 81 cm^2.")],
        ),
        Chapter(
            "Statistika, Peluang, dan Kombinatorika",
            "Data tidak hanya dihitung, tetapi juga ditafsirkan. Peserta harus tahu kapan memakai rata-rata, median, modus, peluang, dan kombinasi.",
            [
                ("Ukuran pemusatan", "Mean sensitif terhadap pencilan, sedangkan median lebih stabil pada data ekstrem."),
                ("Peluang", "Peluang = kejadian yang diinginkan / semua kemungkinan, jika semua kemungkinan setara."),
                ("Kombinasi", "Kombinasi dipakai saat urutan tidak diperhatikan, misalnya memilih panitia."),
            ],
            ["Peluang dua dadu", "Memilih objek dari beberapa kelompok", "Median pada data berpencilan"],
            ex("Data 4,5,5,6,30. Ukuran yang paling mewakili adalah ...", "B", "Ada pencilan 30, sehingga median lebih stabil.", ["A. Mean", "B. Median", "C. Jangkauan", "D. Variansi", "E. Semua sama"]),
            [ex("Peluang muncul angka genap pada satu dadu adalah ...", "C", "Genap: 2,4,6, tiga dari enam = 1/2."), ex("Banyak cara memilih 2 dari 5 orang adalah ...", "A", "C(5,2)=10.")],
        ),
    ]

    reading_chapters = [
        Chapter(
            "Gagasan Utama dan Struktur Paragraf",
            "Literasi membaca dimulai dari kemampuan menemukan gagasan pusat. Tanpa ini, peserta mudah tersesat oleh detail yang sengaja dibuat menarik.",
            [
                ("Gagasan utama", "Gagasan utama adalah ide yang menaungi seluruh paragraf. Ia bisa berada di awal, akhir, atau tersirat dari keseluruhan isi."),
                ("Kalimat penjelas", "Kalimat penjelas berfungsi memberi bukti, contoh, data, atau penjabaran. Detail tidak boleh dipilih sebagai ide utama jika cakupannya terlalu sempit."),
                ("Kohesi", "Kata rujukan seperti ini, tersebut, namun, karena itu, dan sementara itu menunjukkan hubungan antarbagian."),
            ],
            ["Menentukan judul tepat", "Menentukan kalimat utama", "Menyusun paragraf acak"],
            ex("Paragraf membahas penyebab rendahnya minat baca dan menawarkan solusi sekolah. Judul paling tepat adalah ...", "B", "Judul harus mencakup masalah dan solusi, tidak hanya salah satu detail.", ["A. Perpustakaan Sekolah", "B. Strategi Meningkatkan Minat Baca", "C. Buku Digital Mahal", "D. Guru Bahasa", "E. Kebiasaan Siswa"]),
            [ex("Kalimat utama biasanya dapat diuji dengan cara ...", "A", "Kalimat lain dapat menjadi penjelasnya."), ex("Kata 'namun' menunjukkan hubungan ...", "B", "Namun menandai pertentangan.")],
        ),
        Chapter(
            "Inferensi dan Makna Tersirat",
            "Inferensi adalah kesimpulan yang tidak tertulis langsung tetapi wajib mengikuti isi teks. Ini inti soal HOTS literasi.",
            [
                ("Batas inferensi", "Inferensi tidak boleh melompat terlalu jauh. Jawaban harus masih bisa ditopang oleh bukti dalam teks."),
                ("Nada penulis", "Kata evaluatif seperti mengkhawatirkan, penting, keliru, dan mendesak menunjukkan sikap penulis."),
                ("Implikasi", "Implikasi adalah akibat logis jika isi teks diterima sebagai benar."),
            ],
            ["Simpulan tersirat", "Sikap penulis", "Pernyataan yang sesuai/tidak sesuai teks"],
            ex("Teks menyatakan bahwa kebijakan tanpa data sering gagal mencapai sasaran. Inferensi tepatnya adalah ...", "C", "Data diperlukan agar kebijakan lebih tepat sasaran.", ["A. Semua kebijakan gagal", "B. Data tidak penting", "C. Data membantu ketepatan kebijakan", "D. Kebijakan harus dihapus", "E. Sasaran tidak perlu"]),
            [ex("Inferensi yang baik harus ...", "A", "Tetap didukung teks."), ex("Nada kritis berarti penulis ...", "D", "Menilai dan menunjukkan kelemahan.")],
        ),
        Chapter(
            "Literasi Bahasa Inggris",
            "Soal bahasa Inggris biasanya menguji reading comprehension, vocabulary in context, reference, dan inference. Fokus utamanya tetap pemahaman teks.",
            [
                ("Skimming dan scanning", "Skimming untuk gambaran umum, scanning untuk mencari data spesifik seperti angka, tahun, nama, atau istilah."),
                ("Vocabulary in context", "Makna kata ditentukan oleh konteks kalimat, bukan hanya arti kamus yang pertama muncul."),
                ("Reference", "Kata it, they, this, dan those harus dilacak ke nomina yang paling dekat dan paling masuk akal."),
            ],
            ["Main idea", "Synonym in context", "Pronoun reference", "Author's purpose"],
            ex("In a passage, 'this policy' usually refers to ...", "A", "Rujukan biasanya kembali ke kebijakan yang disebut tepat sebelumnya.", ["A. The policy mentioned before", "B. A random future policy", "C. The reader's opinion", "D. The title only", "E. None"]),
            [ex("The word 'decline' in economic text often means ...", "B", "Decline berarti penurunan."), ex("Author's purpose asks about ...", "C", "Tujuan penulis: inform, persuade, explain, criticize.")],
        ),
        Chapter(
            "Literasi Bahasa Arab Dasar",
            "Untuk literasi Arab, peserta perlu mengenali kosakata umum, struktur jumlah ismiyyah/fi'liyyah, dan petunjuk konteks. Tidak semua soal menuntut nahwu berat, tetapi dasar struktur sangat membantu.",
            [
                ("Kosakata inti", "Kelompok kata tentang sekolah, ibadah, keluarga, waktu, tempat, dan akhlak sering menjadi dasar bacaan sederhana."),
                ("Jumlah ismiyyah", "Kalimat nominal umumnya diawali isim. Polanya membantu menemukan subjek dan informasi tentang subjek."),
                ("Jumlah fi'liyyah", "Kalimat verbal diawali fi'il. Perhatikan pelaku, objek, dan keterangan waktu/tempat."),
            ],
            ["Menentukan makna kata", "Menentukan isi bacaan pendek", "Mencari rujukan dhamir"],
            ex("Dalam bacaan Arab sederhana, kata yang berulang di awal beberapa kalimat sering menjadi ...", "A", "Kata berulang bisa menjadi topik utama bacaan.", ["A. Topik", "B. Opsi salah", "C. Tanda baca", "D. Terjemahan", "E. Angka"]),
            [ex("Dhamir perlu dilacak agar pembaca tahu ...", "C", "Dhamir menunjukkan rujukan orang/benda."), ex("Fi'il biasanya menunjukkan ...", "B", "Fi'il menunjukkan tindakan/peristiwa.")],
        ),
    ]

    islam_chapters = [
        Chapter(
            "Quran dan Hadits sebagai Sumber Ajaran",
            "Bab ini menekankan fungsi Quran dan Hadits dalam membimbing akidah, ibadah, akhlak, dan kehidupan sosial. Soal HOTS biasanya meminta penerapan nilai pada kasus.",
            [
                ("Quran", "Quran menjadi sumber utama ajaran Islam. Dalam soal, peserta sering diminta memahami pesan umum ayat, bukan sekadar menghafal potongan teks."),
                ("Hadits", "Hadits menjelaskan, merinci, dan mencontohkan penerapan ajaran. Pahami tema hadits seperti niat, amanah, adab, ilmu, dan kejujuran."),
                ("Kontekstualisasi", "Nilai umum seperti keadilan, kasih sayang, dan tanggung jawab perlu diterapkan pada masalah modern secara proporsional."),
            ],
            ["Penerapan nilai ayat/hadits", "Fungsi hadits terhadap Quran", "Tema akhlak dalam kasus sosial"],
            ex("Seorang siswa menemukan dompet di kelas dan menyerahkannya kepada guru. Nilai utama yang tampak adalah ...", "C", "Menjaga amanah dan kejujuran.", ["A. Riya", "B. Boros", "C. Amanah", "D. Takabur", "E. Ghibah"]),
            [ex("Hadits berfungsi antara lain untuk ...", "B", "Hadits menjelaskan dan mencontohkan ajaran."), ex("Penerapan ajaran harus memperhatikan ...", "A", "Dalil, tujuan, dan konteks.")],
        ),
        Chapter(
            "Aqidah dan Akhlak",
            "Aqidah memberi dasar keyakinan, sedangkan akhlak menunjukkan buah keyakinan dalam tindakan. Keduanya sering muncul dalam soal kasus.",
            [
                ("Iman dan konsekuensi", "Iman tidak berhenti pada pengakuan, tetapi mendorong tanggung jawab moral."),
                ("Akhlak terpuji", "Jujur, sabar, tawadhu, amanah, dan adil sering diuji melalui situasi kehidupan sekolah, keluarga, dan masyarakat."),
                ("Akhlak tercela", "Ghibah, fitnah, hasad, riya, dan takabur perlu dikenali bukan hanya definisinya, tetapi dampaknya."),
            ],
            ["Membedakan akhlak terpuji dan tercela", "Menentukan sikap sesuai nilai iman", "Analisis kasus pergaulan digital"],
            ex("Menyebarkan kabar buruk yang belum terbukti tentang teman termasuk ...", "B", "Menyebarkan kabar tanpa bukti dekat dengan fitnah dan merusak kehormatan.", ["A. Amanah", "B. Fitnah", "C. Tawadhu", "D. Ikhlas", "E. Syukur"]),
            [ex("Sikap tawadhu berarti ...", "A", "Rendah hati tanpa merendahkan diri."), ex("Riya berbahaya karena ...", "C", "Menggeser niat ibadah kepada pujian manusia.")],
        ),
        Chapter(
            "Fikih Ibadah dan Muamalah",
            "Fikih mengatur tata cara ibadah dan hubungan sosial. Soal yang baik tidak hanya bertanya hukum, tetapi alasan dan tujuan syariat.",
            [
                ("Ibadah mahdhah", "Shalat, puasa, zakat, dan haji memiliki aturan yang relatif baku. Pelajari syarat, rukun, sunnah, dan hal yang membatalkan."),
                ("Muamalah", "Muamalah mengatur transaksi dan hubungan sosial. Prinsip pentingnya adalah kerelaan, keadilan, kejelasan akad, dan menghindari zalim."),
                ("Maqashid", "Tujuan syariat mencakup penjagaan agama, jiwa, akal, keturunan, dan harta. Ini membantu menjawab soal konteks modern."),
            ],
            ["Syarat dan rukun ibadah", "Etika transaksi", "Maqashid syariah dalam kasus"],
            ex("Transaksi yang informasinya sengaja disembunyikan bertentangan dengan prinsip ...", "D", "Muamalah menuntut kejelasan dan kejujuran.", ["A. Hiburan", "B. Kebiasaan", "C. Persaingan", "D. Kejelasan akad", "E. Perjalanan"]),
            [ex("Zakat memiliki dimensi sosial karena ...", "B", "Zakat membantu distribusi dan kepedulian."), ex("Menjaga akal dalam maqashid berkaitan dengan larangan ...", "A", "Hal yang merusak akal.")],
        ),
        Chapter(
            "Sejarah Kebudayaan Islam dan Moderasi",
            "SKI bukan hafalan tahun semata. Peserta perlu melihat sebab, tokoh, kontribusi ilmu, dan nilai peradaban. Moderasi beragama menguji sikap adil dan seimbang.",
            [
                ("Peradaban ilmu", "Ilmuwan Muslim berkontribusi dalam matematika, kedokteran, astronomi, filsafat, dan pendidikan. Pahami bidang dan dampaknya."),
                ("Islam di Indonesia", "Penyebaran Islam di Nusantara berlangsung melalui perdagangan, pendidikan, dakwah kultural, dan jaringan ulama."),
                ("Moderasi", "Moderasi bukan mencampuradukkan akidah, tetapi bersikap adil, tidak ekstrem, menghargai kemanusiaan, dan menjaga prinsip ajaran."),
            ],
            ["Kontribusi ilmuwan Muslim", "Organisasi Islam Indonesia", "Sikap moderat dalam masyarakat majemuk"],
            ex("Al-Khawarizmi dikenal penting dalam perkembangan ...", "B", "Karyanya menjadi dasar aljabar dan istilah algoritma.", ["A. Seni lukis", "B. Aljabar", "C. Musik", "D. Geologi", "E. Sastra modern"]),
            [ex("Dakwah kultural di Nusantara menunjukkan ...", "C", "Ajaran disampaikan dengan memperhatikan budaya tanpa menghilangkan prinsip."), ex("Moderasi beragama menolak ...", "A", "Sikap ekstrem dan kekerasan.")],
        ),
    ]

    return [
        Module(
            "00",
            "Roadmap Bimbel UM-PTKIN A-Z",
            "Strategi belajar, jadwal, evaluasi, dan cara memakai semua modul",
            "roadmap",
            "Modul pembuka ini berfungsi sebagai buku pegangan tutor dan siswa. Isinya mengatur arah belajar dari nol: mengenal bentuk ujian, membangun rutinitas, mengelola waktu, mencatat kesalahan, dan menyiapkan simulasi.",
            roadmap_chapters,
            [("HOTS", "Kemampuan berpikir tingkat tinggi: analisis, evaluasi, dan simpulan."), ("Diagnosis", "Tes awal untuk mengetahui kekuatan dan kelemahan."), ("Debrief", "Pembahasan setelah latihan untuk menemukan pola salah.")],
        ),
        Module(
            "01",
            "Penalaran Akademik",
            "Logika, verbal, pola, argumen, dan strategi inferensi",
            "logic",
            "Modul ini melatih cara berpikir sistematis. Fokusnya adalah membaca data, menemukan relasi, menilai argumen, dan memilih jawaban yang paling sah secara logis.",
            pa_chapters,
            [("Premis", "Pernyataan dasar dalam argumen."), ("Inferensi", "Kesimpulan yang ditarik dari data."), ("Kuantor", "Penanda jumlah seperti semua, sebagian, tidak semua.")],
        ),
        Module(
            "02",
            "Penalaran Matematika",
            "Aljabar, fungsi, geometri, statistika, peluang, dan model cerita",
            "math",
            "Modul ini mengajarkan matematika sebagai alat bernalar. Rumus dibahas bersama makna, langkah model, dan cara menghindari salah baca.",
            math_chapters,
            [("Variabel", "Lambang untuk nilai yang belum diketahui."), ("Gradien", "Laju perubahan pada garis."), ("Median", "Nilai tengah setelah data diurutkan.")],
        ),
        Module(
            "03",
            "Literasi Membaca",
            "Bahasa Indonesia, Inggris, Arab, inferensi, dan analisis teks",
            "reading",
            "Modul ini membangun kemampuan membaca cepat tetapi akurat. Peserta belajar menemukan gagasan, sikap penulis, makna kata, rujukan, dan simpulan tersirat.",
            reading_chapters,
            [("Gagasan utama", "Ide pusat paragraf atau teks."), ("Kohesi", "Keterkaitan antarkalimat."), ("Reference", "Rujukan kata ganti dalam teks bahasa Inggris.")],
        ),
        Module(
            "04",
            "Literasi Ajaran Islam",
            "Quran-Hadits, aqidah-akhlak, fikih, SKI, dan moderasi",
            "islam",
            "Modul ini membantu siswa memahami ajaran Islam secara konseptual dan aplikatif. Fokusnya bukan hafalan kering, tetapi penerapan nilai pada kasus kehidupan.",
            islam_chapters,
            [("Amanah", "Sikap dapat dipercaya dan bertanggung jawab."), ("Maqashid", "Tujuan dasar syariat."), ("Moderasi", "Sikap adil, seimbang, dan tidak ekstrem.")],
        ),
    ]


def main() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    images = create_images()
    outputs = [write_module(module, images) for module in build_modules()]
    for output in outputs:
        print(output)


if __name__ == "__main__":
    main()
